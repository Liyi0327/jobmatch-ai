import fitz
from docx import Document


def extract_text_from_pdf(uploaded_file) -> str:
    """
    从 PDF 简历中提取文本。
    """
    uploaded_file.seek(0)
    file_bytes = uploaded_file.read()

    text = ""

    with fitz.open(stream=file_bytes, filetype="pdf") as doc:
        for page in doc:
            page_text = page.get_text()
            if page_text:
                text += page_text + "\n"

    return text.strip()


def extract_text_from_docx(uploaded_file) -> str:
    """
    从 Word .docx 简历中提取文本。
    """
    uploaded_file.seek(0)

    document = Document(uploaded_file)

    paragraphs = []

    for paragraph in document.paragraphs:
        content = paragraph.text.strip()
        if content:
            paragraphs.append(content)

    for table in document.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_text.append(cell_text)
            if row_text:
                paragraphs.append(" | ".join(row_text))

    return "\n".join(paragraphs).strip()


def extract_text_from_resume(uploaded_file) -> str:
    """
    根据文件类型自动解析简历。
    当前支持：
    1. PDF
    2. DOCX
    """
    file_name = uploaded_file.name.lower()

    if file_name.endswith(".pdf"):
        return extract_text_from_pdf(uploaded_file)

    if file_name.endswith(".docx"):
        return extract_text_from_docx(uploaded_file)

    raise ValueError("当前仅支持 PDF 或 DOCX 格式简历。")